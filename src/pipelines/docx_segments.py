from __future__ import annotations
from dataclasses import dataclass
from typing import List, Any, Tuple
import docx


@dataclass
class Segment:
    seg_id: str
    kind: str          # "P" paragraph / "C" cell / "H" header / "F" footer
    obj: Any           # python-docx object reference
    text: str
    global_start: int
    global_end: int


def extract_docx_segments(
    path: str,
) -> Tuple[str, List[Segment], List[str]]:
    warnings = []
    d = docx.Document(path)

    # Track changes detection
    xml = d.element.xml
    if "<w:ins" in xml or "<w:del" in xml:
        warnings.append("warning:track_changes_detected")

    segments: List[Segment] = []
    buf = []
    cur = 0
    idx = 0

    def add_seg(kind: str, obj: Any, text: str):
        nonlocal cur, idx
        t = text or ""
        seg_id = f"{kind}{idx:06d}"
        start = cur
        end = cur + len(t)
        segments.append(
            Segment(seg_id, kind, obj, t, start, end)
        )
        buf.append(t)
        buf.append("\n")
        cur = end + 1
        idx += 1

    # headers / footers
    for sec in d.sections:
        try:
            if sec.header and sec.header.paragraphs:
                add_seg(
                    "H", sec.header, sec.header.paragraphs[0].text
                )
        except Exception:
            pass
        try:
            if sec.footer and sec.footer.paragraphs:
                add_seg(
                    "F", sec.footer, sec.footer.paragraphs[0].text
                )
        except Exception:
            pass

    # paragraphs
    for p in d.paragraphs:
        add_seg("P", p, p.text)

    # tables
    for tb in d.tables:
        for row in tb.rows:
            for cell in row.cells:
                add_seg("C", cell, cell.text)

    full_text = "".join(buf)
    return full_text, segments, warnings
