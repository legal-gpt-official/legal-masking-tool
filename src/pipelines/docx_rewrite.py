"""DOCX rewrite with surgical Run-splitting replacement.

Strategy (improved):
  When a replacement spans multiple Runs with different formatting:

  BEFORE (v1.0):
    Run1[bold:"Hello "] + Run2[italic:"World"] → replace "lo Wor" →
    Run1[bold:"Helld"] + Run2[italic:""]  ← italic format LOST

  AFTER (v1.1 — Run splitting):
    Run1[bold:"Hello "] + Run2[italic:"World"] → replace "lo Wor" →
    Run1a[bold:"Hel"] + Run1b[bold:"MASKED"] + Run2a[italic:"ld"]
    ↑ Split Run1 at offset 3, keep formatting on both halves.
    ↑ Split Run2 at offset 3, keep formatting on suffix.

  Implementation:
    1. Build run→offset map
    2. For each replacement (reverse order):
       a. Split the first affected run at the local start offset
       b. Split the last affected run at the local end offset
       c. Set the replacement text on the "consumed" portion of first run
       d. Clear all fully-consumed middle runs
       e. Suffix portion of last run retains its original formatting
"""

from __future__ import annotations

from typing import Any, Dict, Iterator, List, Tuple
import docx
from copy import deepcopy
from lxml import etree


# ---------------------------------------------------------------------------
# Run splitting (OpenXML level)
# ---------------------------------------------------------------------------


def _split_run_at(run, offset: int):
    """Split a Run into two at the given character offset.

    Returns (run_before, run_after) where:
      - run_before has text[:offset] with original formatting
      - run_after has text[offset:] with identical formatting

    The split is done at the OpenXML level to preserve ALL formatting
    properties (font, size, bold, italic, color, underline, highlight, etc.)
    """
    text = run.text or ""
    if offset <= 0:
        return None, run
    if offset >= len(text):
        return run, None

    before_text = text[:offset]
    after_text = text[offset:]

    # Create a deep copy of the run's XML element for the "after" portion
    run_elem = run._element
    new_elem = deepcopy(run_elem)

    # Insert the new element right after the original
    run_elem.addnext(new_elem)

    # Set texts
    run.text = before_text

    # The new element is now in the paragraph; get it as a Run
    # We access it through the paragraph
    from docx.text.run import Run
    new_run = Run(new_elem, run._element.getparent())
    new_run.text = after_text

    return run, new_run


def _apply_backward(text: str, repls: List[Dict[str, Any]]) -> str:
    """Backward text replacement (fallback)."""
    out = text
    for r in sorted(repls, key=lambda x: x["start"], reverse=True):
        s, e = int(r["start"]), int(r["end"])
        out = out[:s] + str(r["replacement"]) + out[e:]
    return out


# ---------------------------------------------------------------------------
# Core replacement logic
# ---------------------------------------------------------------------------


def _apply_replacements_to_runs(
    runs: list,
    replacements: List[Dict[str, Any]],
) -> None:
    """Apply text replacements to a list of Run objects with run splitting."""
    if not runs or not replacements:
        return

    # Build run→offset map (refreshed for each replacement pass)
    def build_map():
        info = []
        pos = 0
        for r in runs:
            t = r.text or ""
            info.append({"run": r, "start": pos, "end": pos + len(t)})
            pos += len(t)
        return info

    # Apply in reverse order so earlier offsets stay valid
    for repl in sorted(replacements, key=lambda r: r["start"], reverse=True):
        rs = int(repl["start"])
        re_ = int(repl["end"])
        new_text = str(repl.get("replacement", ""))

        run_info = build_map()

        # Find affected run indices
        affected = [
            i for i, ri in enumerate(run_info)
            if ri["start"] < re_ and ri["end"] > rs
        ]

        if not affected:
            continue

        if len(affected) == 1:
            # --- Single run: simple in-place text edit ---
            idx = affected[0]
            ri = run_info[idx]
            local_s = rs - ri["start"]
            local_e = re_ - ri["start"]
            old = ri["run"].text or ""
            ri["run"].text = old[:local_s] + new_text + old[local_e:]

        else:
            # --- Multi-run with run splitting ---
            first_idx = affected[0]
            last_idx = affected[-1]
            first_ri = run_info[first_idx]
            last_ri = run_info[last_idx]

            # Calculate local offsets within first and last runs
            local_s_first = rs - first_ri["start"]
            local_e_last = re_ - last_ri["start"]

            # Split last run at the end offset (if not at boundary)
            last_run = last_ri["run"]
            last_text = last_run.text or ""
            suffix_text = last_text[local_e_last:]

            if local_e_last < len(last_text) and suffix_text:
                # There's text after the replacement in the last run
                # Split to preserve its formatting
                try:
                    _, after_run = _split_run_at(last_run, local_e_last)
                    # Clear the consumed portion of last_run
                    last_run.text = ""
                except Exception:
                    # Fallback: just set the text
                    last_run.text = suffix_text
                    suffix_text = ""
            else:
                # Last run is fully consumed
                last_run.text = ""

            # Set first run: prefix + replacement
            first_run = first_ri["run"]
            first_text = first_run.text or ""
            prefix = first_text[:local_s_first]
            first_run.text = prefix + new_text

            # Clear all middle runs (fully consumed)
            for idx in affected[1:-1]:
                run_info[idx]["run"].text = ""

            # If last run was fully consumed (no split happened), clear it
            if not suffix_text and last_idx != first_idx:
                last_run.text = ""

        # Rebuild runs list after splits may have added elements
        _refresh_runs_list(runs)


def _refresh_runs_list(runs: list) -> None:
    """After run splitting, the paragraph may have new Run elements.
    This is a no-op since we work on the actual Run objects in-place.
    python-docx's paragraph.runs will reflect splits automatically.
    """
    pass


def _apply_replacements_to_paragraph(
    p,
    replacements: List[Dict[str, Any]],
) -> None:
    """Apply replacements to a paragraph, preferring run-level edits."""
    runs = list(p.runs)
    if runs:
        _apply_replacements_to_runs(runs, replacements)
    else:
        old_text = p.text or ""
        new_text = _apply_backward(old_text, replacements)
        p.text = ""
        p.add_run(new_text)


def _apply_replacements_to_cell(
    cell,
    replacements: List[Dict[str, Any]],
) -> None:
    """Apply replacements to a table cell."""
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        return

    if len(paragraphs) == 1:
        _apply_replacements_to_paragraph(paragraphs[0], replacements)
        return

    # Multi-paragraph cell: build paragraph offset map
    para_map: List[Dict[str, Any]] = []
    pos = 0
    for p in paragraphs:
        text = p.text or ""
        para_map.append({"para": p, "start": pos, "end": pos + len(text)})
        pos += len(text) + 1  # +1 for \n separator

    for repl in sorted(replacements, key=lambda r: r["start"], reverse=True):
        rs = int(repl["start"])
        re_ = int(repl["end"])

        target = None
        for pi in para_map:
            if rs >= pi["start"] and re_ <= pi["end"]:
                target = pi
                break

        if target is not None:
            local_repl = {
                "start": rs - target["start"],
                "end": re_ - target["start"],
                "replacement": repl["replacement"],
            }
            _apply_replacements_to_paragraph(target["para"], [local_repl])
        else:
            # Cross-paragraph (rare): fallback
            full_text = cell.text or ""
            new_text = full_text[:rs] + str(repl["replacement"]) + full_text[re_:]
            if paragraphs:
                _apply_replacements_to_paragraph(
                    paragraphs[0],
                    [{"start": 0, "end": len(paragraphs[0].text or ""),
                      "replacement": new_text}],
                )
                for extra in paragraphs[1:]:
                    for run in list(extra.runs):
                        run.text = ""


# ---------------------------------------------------------------------------
# Segment iterator
# ---------------------------------------------------------------------------


def _iter_segments_same_doc(
    d: docx.Document,
) -> Iterator[Tuple[str, str, Any, str]]:
    """Yield (seg_id, kind, obj, text) matching docx_segments ordering."""
    idx = 0

    def mk(kind: str) -> str:
        nonlocal idx
        seg_id = f"{kind}{idx:06d}"
        idx += 1
        return seg_id

    for sec in d.sections:
        try:
            if sec.header and sec.header.paragraphs:
                p = sec.header.paragraphs[0]
                yield mk("H"), "H", p, p.text or ""
        except Exception:
            pass
        try:
            if sec.footer and sec.footer.paragraphs:
                p = sec.footer.paragraphs[0]
                yield mk("F"), "F", p, p.text or ""
        except Exception:
            pass

    for p in d.paragraphs:
        yield mk("P"), "P", p, p.text or ""

    for tb in d.tables:
        for row in tb.rows:
            for cell in row.cells:
                yield mk("C"), "C", cell, cell.text or ""


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def rewrite_docx_with_maps(
    src_path: str,
    dst_path: str,
    docx_maps: List[Dict[str, Any]],
) -> None:
    """Rewrite DOCX with format-preserving run-splitting replacements."""
    d = docx.Document(src_path)

    by_seg: Dict[str, List[Dict[str, Any]]] = {}
    for m in docx_maps or []:
        sid = m.get("seg_id")
        if not sid:
            continue
        by_seg.setdefault(sid, []).append(m)

    for seg_id, kind, obj, text in _iter_segments_same_doc(d):
        reps = by_seg.get(seg_id)
        if not reps:
            continue

        local = [
            {
                "start": r.get("local_start", 0),
                "end": r.get("local_end", 0),
                "replacement": r.get("replacement", ""),
            }
            for r in reps
        ]

        if kind == "P":
            _apply_replacements_to_paragraph(obj, local)
        elif kind == "C":
            _apply_replacements_to_cell(obj, local)
        elif kind in ("H", "F"):
            try:
                _apply_replacements_to_paragraph(obj, local)
            except Exception:
                new_text = _apply_backward(text, local)
                try:
                    obj.text = new_text
                except Exception:
                    pass

    d.save(dst_path)
